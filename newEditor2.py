import os
import subprocess
import time

import mammoth
import pypandoc
from PyQt5 import QtPrintSupport, QtGui
from PyQt5.QtCore import Qt, QTimer
from PyQt5.QtGui import QFont, QTextListFormat, QTextCharFormat, QTextCursor, QTextDocument, \
    QTextBlockFormat
from PyQt5.QtGui import QIcon
from PyQt5.QtPrintSupport import QPrinter
from PyQt5.QtWidgets import (QApplication, QDialog, QMainWindow, QScrollArea,
                             QFileDialog, QSizePolicy, QMessageBox, QColorDialog)
from PyQt5.QtWidgets import QWidget, QVBoxLayout
from docx import Document

from editor.actions.editor_actions import EditorActions
from editor.common_Dialogs import CommonDialogs
from editor.components.ascii_unicode_ConversionDialog import ConversionDialog
from editor.components.customize_image import ImageEditDialog
from editor.components.excel_csv_file_handling import ExcelCsvViewer
from editor.components.format_content import SpacingDialog
from editor.components.new_editor_components import NewPageLayoutDialog, NewPage
from editor.components.speech_to_text import LanguageSelectionPopup, SpeechToTextThread
from editor.widgets.zoom_slider import ZoomSlider
from logger import setup_logger
from spellcheck.bloom_filter import start_bloom
from utils.asciitounicode import process_line
from utils.corpus_clean import get_clean_words_for_dictionary
from utils.find import Find
from utils.sort_by import SortDialog
from utils.table import Table
from utils.wordcount import WordCount

filename = os.path.splitext(os.path.basename(__file__))[0]

# Set up logger
logger = setup_logger(filename)


def start_background_exe():
    exe_path = r"resources\keyboardDriver\kannadaKeyboard.exe"  # Path to your executable relative to the current directory
    # exe_path = r"resources\keyboardDriver\testing.exe"
    try:
        # Use subprocess.Popen to start the executable in the background
        print("kannada Nudi Keyboard loaded and running in background")
        subprocess.Popen([exe_path], shell=True, stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                         start_new_session=True)

    except Exception as e:
        print(f"Error starting background exe: {e}")


class NewTextEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.statusbar = None
        self.scroll_layout = None
        self.scroll_content = None
        self.scroll_area = None
        self.actions = EditorActions(self)
        self.current_file_path = None
        self.total_pages = 0
        self.pages = []
        self.current_page = None
        self.filename = None
        self.speech_thread = None
        self.editor_windows = []  # Add this line to keep references to new editor windows
        self.error_dialog = CommonDialogs()
        self.initUI()
        # Initialize and add ZoomSlider
        self.zoom_slider = ZoomSlider()
        self.zoom_slider.initZoomSlider(self)
    def initUI(self):
        # start_background_exe()
        self.actions.createActions()
        self.actions.createMenus()
        self.actions.createToolbars()
        self.actions.createFormatbar()

        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        layout = QVBoxLayout(central_widget)
        layout.setContentsMargins(0, 0, 0, 0)  # Remove outer margins
        layout.setSpacing(0)  # Remove spacing between elements

        self.scroll_area = QScrollArea(self)
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.scroll_content = QWidget()
        self.scroll_layout = QVBoxLayout(self.scroll_content)
        self.scroll_layout.setContentsMargins(20, 20, 20, 20)  # Adjust margins as needed
        self.scroll_layout.setAlignment(Qt.AlignTop | Qt.AlignHCenter)
        self.scroll_area.setWidget(self.scroll_content)

        layout.addWidget(self.scroll_area)

        self.statusbar = self.statusBar()

        self.addNewPage()
        self.current_page.editor.installEventFilter(self)
        self.setGeometry(100, 100, 1030, 800)
        self.setWindowTitle("ಕನ್ನಡ ನುಡಿ - " + self.access_filename())
        self.setWindowIcon(QIcon('resources/images/logo.jpg'))  # Set the application icon
        self.showMaximized()
        self.setFocusToEditor()
        self.statusbar.setStatusTip("total pages: " + str(self.total_pages))

    def setFocusToEditor(self):
        # Ensure focus is set to the editor of the current page
        if self.current_page and hasattr(self.current_page, 'editor'):
            self.current_page.editor.setFocus()

    def updateZoom(self, value):
        factor = value / 100
        for page in self.pages:
            page.setZoomFactor(factor)

    def addNewPage(self):
        page = NewPage(self)
        page.textOverflow.connect(self.handleTextOverflow)  # Connect the textOverflow signal
        page.clicked.connect(self.setActivePage)
        self.pages.append(page)
        self.scroll_layout.addWidget(page)
        self.setActivePage(page)
        self.total_pages += 1

        # Return the page so we can further customize it if needed
        return page

    def setActivePage(self, page):
        self.current_page = page

    def handleTextOverflow(self):
        if self.current_page:
            # Step 1: Create the new page and get a reference to it
            new_page = self.addNewPage()

            # Step 2: Move the overflowed content to the new page
            remaining_text = self.current_page.editor.toPlainText()

            # Clear the current page's editor
            self.current_page.editor.clear()

            # Set the text in the new page's editor
            new_page.editor.insertPlainText(remaining_text)

            # Step 3: Move the cursor to the new page and set focus there
            new_page.editor.setFocus()  # Focus the editor of the new page
            cursor = new_page.editor.textCursor()
            cursor.movePosition(QtGui.QTextCursor.Start)  # Move cursor to the start of the new page
            new_page.editor.setTextCursor(cursor)

            # Optional: Scroll to ensure the new page is visible
            self.scroll_layout.addWidget(new_page)  # Adjust this based on your scrolling logic

    # Override the closeEvent method
    def closeEvent(self, event):
        try:
            if self.speech_thread:
                self.speech_thread.stop()
            event.accept()
        except Exception as e:
            self.error_dialog.show_error_popup(str(e))

        if any(page.is_changed for page in self.pages):
            reply = QMessageBox.question(self, 'ಸಂದೇಶ',
                                         "ನೀವು ಉಳಿಸದ ಬದಲಾವಣೆಗಳನ್ನು ಹೊಂದಿರುವಿರಿ. ನಿಮ್ಮ ಬದಲಾವಣೆಗಳನ್ನು ನೀವು ಉಳಿಸಲು ಬಯಸುವಿರಾ?",
                                         QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel,
                                         QMessageBox.Cancel)
            if reply == QMessageBox.Yes:
                self.saveFile()
                event.accept()
            elif reply == QMessageBox.No:
                event.accept()
            else:
                event.ignore()
        else:
            event.accept()

    def newFile(self):
        self.pages.clear()
        for i in reversed(range(self.scroll_layout.count())):
            widget = self.scroll_layout.itemAt(i).widget()
            if widget is not None:
                widget.deleteLater()
        self.addNewPage()

    def openFile(self):
        options = QFileDialog.Options()
        self.filename, _ = QFileDialog.getOpenFileName(
            self,
            "Open File",
            "",
            "All Files (*);;Text Files (*.txt);;Word Documents (*.docx);;Rich Text Format (*.rtf)",
            options=options
        )
        if self.filename:
            content = ""
            file_extension = self.filename.split('.')[-1].lower()

            try:
                if file_extension == 'txt':
                    with open(self.filename, 'r', encoding="utf-8") as file:
                        content = file.read()
                elif file_extension == 'docx':
                    with open(self.filename, "rb") as docx_file:
                        result = mammoth.convert_to_html(docx_file)
                        content = result.value  # The converted HTML
                elif file_extension == 'rtf':
                    content = pypandoc.convert_file(self.filename, 'plain', format='rtf')
                else:
                    raise ValueError("Unsupported file format")

                # Split content into chunks of approximately 490 words
                words = content.split()
                word_limit = 490
                current_word_count = 0
                current_page_content = []
                for word in words:
                    current_page_content.append(word)
                    current_word_count += 1
                    if current_word_count >= word_limit:
                        self.addPageWithContent(' '.join(current_page_content))
                        current_page_content = []
                        current_word_count = 0

                # Add any remaining content to a new page if not empty
                if current_page_content:
                    self.addPageWithContent(' '.join(current_page_content))

            except Exception as e:
                self.error_dialog.show_error_popup(str(e))

            # Update window title and remove blank pages
            self.setWindowTitle("ಕನ್ನಡ ನುಡಿ - " + self.access_filename())
            self.removeBlankPages()
            self.statusbar.setStatusTip("total pages: " + str(self.total_pages))

            # Set the current file path
            self.current_file_path = self.filename

    def saveFile(self):
        if not self.current_file_path:
            options = QFileDialog.Options()
            self.current_file_path, _ = QFileDialog.getSaveFileName(
                self, "Save File", "",
                "Text Files (*.txt);;Word Files (*.docx);;Rich Text Format (*.rtf);;PDF Files (*.pdf);;All Files (*)",
                options=options
            )

        if self.current_file_path:
            content = self.pages[0].editor.toPlainText()
            try:
                if self.current_file_path.endswith('.txt'):
                    with open(self.current_file_path, 'w', encoding='utf-8') as file:
                        file.write(content)
                elif self.current_file_path.endswith('.docx'):
                    doc = Document()
                    doc.add_paragraph(content)
                    doc.save(self.current_file_path)
                elif self.current_file_path.endswith('.rtf'):
                    pypandoc.convert_text(content, 'rtf', format='md', outputfile=self.current_file_path,
                                          encoding='utf-8')
                elif self.current_file_path.endswith('.pdf'):
                    printer = QPrinter(QPrinter.HighResolution)
                    printer.setOutputFormat(QPrinter.PdfFormat)
                    printer.setOutputFileName(self.current_file_path)
                    doc = QTextDocument()
                    doc.setPlainText(content)
                    doc.print_(printer)
                else:
                    self.error_dialog.showError("Unsupported file format")
            except Exception as e:
                self.error_dialog.show_error_popup(str(e))

    def saveAsFile(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save File", "",
            "Text Files (*.txt);;Word Files (*.docx);;Rich Text Format (*.rtf);;PDF Files (*.pdf);;All Files (*)",
            options=options
        )

        if file_path:
            content = self.pages[0].editor.toPlainText()
            try:
                if file_path.endswith('.txt'):
                    with open(file_path, 'w', encoding='utf-8') as file:
                        file.write(content)
                elif file_path.endswith('.docx'):
                    doc = Document()
                    doc.add_paragraph(content)
                    doc.save(file_path)
                elif file_path.endswith('.rtf'):
                    pypandoc.convert_text(content, 'rtf', format='md', outputfile=file_path, encoding='utf-8')
                elif file_path.endswith('.pdf'):
                    printer = QPrinter(QPrinter.HighResolution)
                    printer.setOutputFormat(QPrinter.PdfFormat)
                    printer.setOutputFileName(file_path)
                    doc = QTextDocument()
                    doc.setPlainText(content)
                    doc.print_(printer)
                else:
                    self.error_dialog.showError("Unsupported file format")
            except Exception as e:
                self.error_dialog.show_error_popup(str(e))

    def openAsciiFile(self):
        options = QFileDialog.Options()
        self.filename, _ = QFileDialog.getOpenFileName(
            self, "Open File", "",
            "All Files (*);;Text Files (*.txt);;Word Documents (*.docx);;Rich Text Format (*.rtf)",
            options=options)
        if self.filename:
            content = ""
            file_extension = self.filename.split('.')[-1].lower()

            try:
                if file_extension == 'txt':
                    with open(self.filename, 'r', encoding="utf-8") as file:
                        unicode_lines = [process_line(line) for line in file]
                        content = ''.join(unicode_lines)
                elif file_extension == 'docx':
                    doc = Document(self.filename)
                    paragraphs = [process_line(para.text) for para in doc.paragraphs]
                    content = "\n".join(paragraphs)
                elif file_extension == 'rtf':
                    plain_text = pypandoc.convert_file(self.filename, 'plain', format='rtf')
                    unicode_lines = [process_line(line) for line in plain_text.splitlines()]
                    content = '\n'.join(unicode_lines)
                else:
                    self.error_dialog("Unsupported file format")
                    # raise ValueError("Unsupported file format")

                # Split content into chunks of approximately 490 words
                words = content.split()
                word_limit = 490
                current_word_count = 0
                current_page_content = []
                for word in words:
                    current_page_content.append(word)
                    current_word_count += 1
                    if current_word_count >= word_limit:
                        self.addPageWithContent(' '.join(current_page_content))
                        current_page_content = []
                        current_word_count = 0

                # Add any remaining content to a new page if not empty
                if current_page_content:
                    self.addPageWithContent(' '.join(current_page_content))

            except Exception as e:
                self.error_dialog.showError(str(e))

            # Update window title and remove blank pages
            self.setWindowTitle("ಕನ್ನಡ ನುಡಿ - " + self.access_filename())
            self.removeBlankPages()

    def removeBlankPages(self):
        for i in reversed(range(len(self.pages))):
            page = self.pages[i]
            text = page.editor.toPlainText().strip()
            if not text:
                self.scroll_layout.removeWidget(page)
                page.deleteLater()
                self.pages.pop(i)
                self.total_pages -= 1

    def addPageWithContent(self, content):
        if not content:
            return self.addNewPage()
        self.addNewPage()
        self.current_page.editor.setHtml(content)
        self.handleTextOverflow()

    def setFontFamily(self, font):
        if self.current_page:
            self.current_page.editor.setCurrentFont(font)

    def setFontSize(self, index):
        size = int(self.actions.fontSizeComboBox.currentText())
        if self.current_page and hasattr(self.current_page, 'editor'):
            try:
                self.current_page.editor.setFontPointSize(size)
            except RuntimeError as e:
                print(f"Error setting font size: {str(e)}")
                # Handle the error as needed (e.g., log it, notify the user)
        else:
            self.setActivePage(self.current_page)
            print("Error: No current page or editor not available")
            # Handle the error as needed (e.g., log it, notify the user)

    def toggleBold(self):
        if self.current_page:
            editor = self.current_page.editor
            cursor = editor.textCursor()

            # Print the current selected text
            selected_text = cursor.selectedText()
            # print(f"Selected Text: '{selected_text}'")

            # Check the current font weight
            current_weight = editor.fontWeight()

            # Toggle the font weight
            new_weight = QFont.Normal if current_weight == QFont.Bold else QFont.Bold
            editor.setFontWeight(new_weight)

    def toggleItalic(self):
        if self.current_page:
            state = self.current_page.editor.fontItalic()
            self.current_page.editor.setFontItalic(not state)

    def toggleUnderline(self):
        if self.current_page:
            state = self.current_page.editor.fontUnderline()
            self.current_page.editor.setFontUnderline(not state)

    def strike(self):

        # Grab the text's format
        fmt = self.current_page.editor.currentCharFormat()

        # Set the fontStrikeOut property to its opposite
        fmt.setFontStrikeOut(not fmt.fontStrikeOut())

        # And set the next char format
        self.current_page.editor.setCurrentCharFormat(fmt)

    def fontColorChanged(self):

        # Get a color from the text dialog
        color = QColorDialog.getColor()

        # Set it as the new text color
        self.current_page.editor.setTextColor(color)

    def superScript(self):

        # Grab the current format
        fmt = self.current_page.editor.currentCharFormat()

        # And get the vertical alignment property
        align = fmt.verticalAlignment()

        # Toggle the state
        if align == QTextCharFormat.AlignNormal:

            fmt.setVerticalAlignment(QTextCharFormat.AlignSuperScript)

        else:

            fmt.setVerticalAlignment(QTextCharFormat.AlignNormal)

        # Set the new format
        self.current_page.editor.setCurrentCharFormat(fmt)

    def subScript(self):

        # Grab the current format
        fmt = self.current_page.editor.currentCharFormat()

        # And get the vertical alignment property
        align = fmt.verticalAlignment()

        # Toggle the state
        if align == QTextCharFormat.AlignNormal:

            fmt.setVerticalAlignment(QTextCharFormat.AlignSubScript)

        else:

            fmt.setVerticalAlignment(QTextCharFormat.AlignNormal)

        # Set the new format
        self.current_page.editor.setCurrentCharFormat(fmt)

    def alignLeft(self):
        self.current_page.editor.setAlignment(Qt.AlignLeft)

    def alignRight(self):
        self.current_page.editor.setAlignment(Qt.AlignRight)

    def alignCenter(self):
        self.current_page.editor.setAlignment(Qt.AlignCenter)

    def alignJustify(self):
        self.current_page.editor.setAlignment(Qt.AlignJustify)

    def indent(self):

        # Grab the cursor
        cursor = self.current_page.editor.textCursor()

        if cursor.hasSelection():

            # Store the current line/block number
            temp = cursor.blockNumber()

            # Move to the selection's end
            cursor.setPosition(cursor.anchor())

            # Calculate range of selection
            diff = cursor.blockNumber() - temp

            direction = QTextCursor.Up if diff > 0 else QTextCursor.Down

            # Iterate over lines (diff absolute value)
            for n in range(abs(diff) + 1):
                # Move to start of each line
                cursor.movePosition(QTextCursor.StartOfLine)

                # Insert tabbing
                cursor.insertText("\t")

                # And move back up
                cursor.movePosition(direction)

        # If there is no selection, just insert a tab
        else:

            cursor.insertText("\t")

    def handleDedent(self, cursor):

        cursor.movePosition(QTextCursor.StartOfLine)

        # Grab the current line
        line = cursor.block().text()

        # If the line starts with a tab character, delete it
        if line.startswith("\t"):

            # Delete next character
            cursor.deleteChar()

        # Otherwise, delete all spaces until a non-space character is met
        else:
            for char in line[:8]:

                if char != " ":
                    break

                cursor.deleteChar()

    def dedent(self):

        cursor = self.current_page.editor.textCursor()

        if cursor.hasSelection():

            # Store the current line/block number
            temp = cursor.blockNumber()

            # Move to the selection's last line
            cursor.setPosition(cursor.anchor())

            # Calculate range of selection
            diff = cursor.blockNumber() - temp

            direction = QTextCursor.Up if diff > 0 else QTextCursor.Down

            # Iterate over lines
            for n in range(abs(diff) + 1):
                self.handleDedent(cursor)

                # Move up
                cursor.movePosition(direction)

        else:
            self.handleDedent(cursor)

    def bulletList(self):

        cursor = self.current_page.editor.textCursor()

        # Insert bulleted list
        cursor.insertList(QTextListFormat.ListDisc)

    def numberList(self):

        cursor = self.current_page.editor.textCursor()

        # Insert list with numbers
        cursor.insertList(QTextListFormat.ListDecimal)

    def highlight(self):

        color = QColorDialog.getColor()

        self.current_page.editor.setTextBackgroundColor(color)

    def setSpacing(self):
        if not self.current_page or not self.current_page.editor:
            print("Error: current_page or editor is not valid.")
            return

        dialog = SpacingDialog(self)
        dialog.setWindowTitle('Line & Paragraph Spacing')

        # Retrieve current font pixel size if available
        current_font = self.current_page.editor.currentFont()
        if current_font:
            dialog.customLineSpacingSpinBox.setValue(current_font.pixelSize())

        dialog.beforeParagraphSpinBox.setValue(0)
        dialog.afterParagraphSpinBox.setValue(0)

        if dialog.exec_() == QDialog.Accepted:
            dialog.applySettings()

    def setLineSpacing(self, value):
        cursor = self.current_page.editor.textCursor()
        fmt = cursor.blockFormat()
        fmt.setLineHeight(value, QTextBlockFormat.LineDistanceHeight)
        cursor.setBlockFormat(fmt)

    # ----------------------------------------------------------------FORMAT functions ----------------------------------------------------------------

    def refresh_recheck(self):
        self.total_pages = 0
        total_words = 0
        total_incorrect_words = 0
        start_time = time.time()

        for page in self.pages:
            self.total_pages += 1
            if not page or not page.editor:
                print("Page or editor is not valid.")
                continue

            # Retrieve plain text from current page's editor
            plain_text = page.editor.toPlainText()

            # Count total words
            words = plain_text.split()
            total_words += len(words)

            # Process text content for spell checking
            content_for_bloom = [get_clean_words_for_dictionary(word) for word in words if len(word) > 1]
            wrong_words = start_bloom(content_for_bloom)

            # Count total incorrect words
            total_incorrect_words += len(wrong_words)

            # Underline incorrect words in the editor
            highlighted_content = plain_text
            for word in wrong_words:
                highlighted_content = highlighted_content.replace(word,
                                                                  f'<span style="text-decoration: underline;">{word}</span>')

            # Update the editor with the underlined content
            page.editor.setHtml(highlighted_content)

        end_time = time.time()
        spellcheck_time = end_time - start_time

        # Show info dialog with the requested information and parent window's logo
        info_msg = QMessageBox()
        info_msg.setWindowTitle("ವರದಿ ಪರೀಕ್ಷೆ ಮಾಹಿತಿ")
        info_msg.setText(f"ಒಟ್ಟು ಪದಗಳ ಸಂಖ್ಯೆ : {total_words}\n"
                         f"ತಪ್ಪು ಪದಗಳ ಸಂಖ್ಯೆ : {total_incorrect_words}\n"
                         f"ಕಾಗುಣಿತ ಪರಿಶೀಲನೆಗಾಗಿ ತೆಗೆದುಕೊಂಡ ಒಟ್ಟು ಸಮಯ : {spellcheck_time:.2f} ಸೆಕೆಂಡುಗಳು")
        info_msg.setIcon(QMessageBox.Information)

        # Set the parent window's icon for the message box
        parent_icon = self.windowIcon()
        if parent_icon:
            info_msg.setWindowIcon(parent_icon)

        info_msg.exec_()

        self.removeBlankPages()
        self.statusBar().showMessage("ಒಟ್ಟು ಪುಟಗಳು: " + str(self.total_pages))

        # Automatically close the message box after 5 seconds
        QTimer.singleShot(5000, info_msg.close)

        # Ensure the application processes the event loop to display the message box
        QApplication.processEvents()

    def zoomIn(self):
        if self.current_page:
            self.current_page.editor.zoomIn()

    def zoomOut(self):
        if self.current_page:
            self.current_page.editor.zoomOut()

    def access_filename(self):
        return self.filename if self.filename is not None else "Untitled"

    def printHandler(self):

        # Open printing dialog
        dialog = QtPrintSupport.QPrintDialog()

        if dialog.exec_() == QDialog.Accepted:
            self.current_page.editor.document().print_(dialog.printer())

    def new(self):
        new_editor = NewTextEditor()
        new_editor.show()
        self.editor_windows.append(new_editor)  # Keep a reference to the new editor window

    def page_layout_size(self):
        dialog = NewPageLayoutDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            width, height = dialog.getPageSize()
            for page in self.pages:
                page.setPageSize(width, height)

    def undo(self):
        if self.current_page:
            self.current_page.editor.undo()

    def redo(self):
        if self.current_page:
            self.current_page.editor.redo()

    def insertTable(self):
        if self.current_page:
            table_dialog = Table(self.current_page.editor)  # Pass the QTextEdit widget to Table
            table_dialog.exec_()

    def find_replace(self):
        if self.current_page:
            find = Find(self.current_page.editor)  # Pass the QTextEdit widget to Table
            find.exec_()

    def choose_image(self):
        file_dialog = QFileDialog(self)
        file_dialog.setNameFilter("Images (*.png *.jpg *.bmp)")
        if file_dialog.exec_():
            file_path = file_dialog.selectedFiles()[0]
            image_dialog = ImageEditDialog(file_path, self)
            if image_dialog.exec_() == QDialog.Accepted:
                modified_image = image_dialog.getModifiedImage()
                if not modified_image.isNull():
                    self.insertEditedImage(modified_image)

    def insertEditedImage(self, edited_image):
        if self.current_page:
            cursor = self.current_page.editor.textCursor()
            cursor.insertImage(edited_image)

    def bulletList(self):

        cursor = self.current_page.editor.textCursor()

        # Insert bulleted list
        cursor.insertList(QTextListFormat.ListDisc)

    def numberList(self):
        if not self.current_page or not self.current_page.editor:
            print("Error: current_page or editor is not valid.")
            return
        cursor = self.current_page.editor.textCursor()

        # Insert list with numbers
        cursor.insertList(QTextListFormat.ListDecimal)

    def sortByAction(self):
        if not self.current_page or not self.current_page.editor:
            print("Error: current_page or editor is not valid.")
            return

        # Create and execute the SortDialog
        sort_dialog = SortDialog(self)
        if sort_dialog.exec_():
            sort_by = sort_dialog.combo_sort_by.currentText()
            type_ = sort_dialog.combo_type.currentText()
            using = sort_dialog.combo_using.currentText()
            ascending = sort_dialog.radio_asc.isChecked()
            has_header = sort_dialog.check_has_header.isChecked()
            separator = sort_dialog.line_separator.text()
            sort_options = sort_dialog.combo_sort_options.currentText()

            # Retrieve text from QTextEdit
            text = self.current_page.editor.toPlainText()
            lines = text.split('\n')

            # Implement sorting based on user input
            if sort_by == "Paragraph":
                # Implement sorting logic for Paragraph sorting
                pass
            elif sort_by == "Headings":
                # Implement sorting logic for Headings sorting
                pass
            elif sort_by == "Fields":
                # Implement sorting logic for Fields sorting
                pass
            elif sort_by == "Header Name":
                # Implement sorting logic for Header Name sorting
                pass
            elif sort_by == "Column Number":
                # Implement sorting logic for Column Number sorting
                pass
            else:
                print("Unknown sort_by option:", sort_by)
                return

            # Example sorting for different types
            if type_ == "Text":
                lines.sort(key=str.lower if not sort_options == "Case sensitive" else str, reverse=not ascending)
            elif type_ == "Number":
                lines.sort(key=lambda x: float(x) if x.replace('.', '', 1).isdigit() else float('inf'),
                           reverse=not ascending)
            elif type_ == "Date":
                from datetime import datetime
                lines.sort(key=lambda x: datetime.strptime(x, '%Y-%m-%d'), reverse=not ascending)

            sorted_text = '\n'.join(lines)

            # Update QTextEdit with sorted text
            self.current_page.editor.setPlainText(sorted_text)

    def toggle_speech_to_text(self):
        sender = self.sender()
        if sender.isChecked():
            popup = LanguageSelectionPopup()
            if popup.exec_() == QDialog.Accepted:
                selected_language = popup.selectedLanguage
                if selected_language is None:
                    QMessageBox.critical(None, 'Error', 'No language selected')
                    sender.setChecked(False)
                    return

                sender.setText("Stop Speech to Text")
                self.speech_thread = SpeechToTextThread(self.current_page.editor, selected_language)
                self.speech_thread.start()
            else:
                sender.setChecked(False)
        else:
            sender.setText("Speech to Text")
            if self.speech_thread:
                self.speech_thread.stop()
                self.speech_thread = None

    def ascii_to_unicode_converter(self):
        dialog = ConversionDialog(self)
        dialog.exec_()

    def excel_csv_file(self):
        self.viewer = ExcelCsvViewer()
        self.viewer.show()

    def wordCount(self):
        wc = WordCount(self.current_page)

        wc.getText()

        wc.show()


if __name__ == "__main__":
    import sys

    app = QApplication(sys.argv)
    editor = NewTextEditor()
    editor.show()
    sys.exit(app.exec_())
