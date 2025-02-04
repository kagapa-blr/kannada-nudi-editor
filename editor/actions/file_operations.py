import pypandoc
from PyQt5.QtGui import QTextDocument
from PyQt5.QtPrintSupport import QPrinter
from PyQt5.QtWidgets import QFileDialog
from docx import Document


class FileOperation:
    def __init__(self, editor):
        self.editor = editor

    def handle_open_file(self):

        options = QFileDialog.Options()
        self.editor.filename, _ = QFileDialog.getOpenFileName(
            self.editor,
            "Open File",
            "",
            "All Files (*);;Text Files (*.txt);;Word Documents (*.docx);;Rich Text Format (*.rtf)",
            options=options
        )
        if self.editor.filename:
            self.editor.content = ""
            file_extension = self.editor.filename.split('.')[-1].lower()

            try:
                if file_extension == 'txt':
                    with open(self.editor.filename, 'r', encoding="utf-8") as file:
                        content = file.read()

                elif file_extension == 'docx':
                    doc = Document(self.editor.filename)
                    content = '\n'.join([para.text for para in doc.paragraphs])

                elif file_extension == 'rtf':
                    content = pypandoc.convert_file(self.editor.filename, 'plain', format='rtf')

                else:
                    raise ValueError("Unsupported file format")

                # Split content into chunks of approximately 490 words
                words = content.split()
                word_limit = 490
                current_word_count = 0
                current_page_content = []

                for word in words:
                    current_page_content.append(word)
                    current_word_count += 1

                    if current_word_count >= word_limit:
                        self.editor.addPageWithContent(' '.join(current_page_content))
                        current_page_content = []
                        current_word_count = 0

                # Add any remaining content to a new page if not empty
                if current_page_content:
                    self.editor.addPageWithContent(' '.join(current_page_content))

            except Exception as e:
                self.editor.error_dialog.show_error_popup(f"Error opening file: {str(e)}")

            # Update window title and remove blank pages
            self.editor.setWindowTitle("ಕನ್ನಡ ನುಡಿ - " + self.editor.access_filename())
            self.editor.removeBlankPages()
            self.editor.statusbar.setStatusTip(f"Total pages: {self.editor.total_pages}")

            # Set the current file path
            self.editor.current_file_path = self.editor.filename


    def handle_save_file(self,content):
        if not self.editor.current_file_path:
            options = QFileDialog.Options()
            self.editor.current_file_path, _ = QFileDialog.getSaveFileName(
                self.editor, "Save File", "",
                "Text Files (*.txt);;Word Files (*.docx);;Rich Text Format (*.rtf);;PDF Files (*.pdf);;All Files (*)",
                options=options
            )

        if self.editor.current_file_path:
            #content = self.editor.get_editor_content()  # Assuming this gets the text from the editor
            print('content', content)
            try:
                if self.editor.current_file_path.endswith('.txt'):
                    with open(self.editor.current_file_path, 'w', encoding='utf-8') as file:
                        file.write(content)
                elif self.editor.current_file_path.endswith('.docx'):
                    doc = Document()
                    doc.add_paragraph(content)
                    doc.save(self.editor.current_file_path)
                elif self.editor.current_file_path.endswith('.rtf'):
                    pypandoc.convert_text(content, 'rtf', format='md', outputfile=self.editor.current_file_path,
                                          encoding='utf-8')
                elif self.editor.current_file_path.endswith('.pdf'):
                    printer = QPrinter(QPrinter.HighResolution)
                    printer.setOutputFormat(QPrinter.PdfFormat)
                    printer.setOutputFileName(self.editor.current_file_path)
                    doc = QTextDocument()
                    doc.setPlainText(content)
                    doc.print_(printer)
                else:
                    self.editor.error_dialog.show_error_popup("Unsupported file format")
            except Exception as e:
                self.editor.error_dialog.show_error_popup(str(e))

    def handle_save_as_file(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            self.editor, "Save File", "",
            "Text Files (*.txt);;Word Files (*.docx);;Rich Text Format (*.rtf);;PDF Files (*.pdf);;All Files (*)",
            options=options
        )

        if file_path:
            content = self.editor.get_current_text()  # Get text from editor
            try:
                if file_path.endswith('.txt'):
                    with open(file_path, 'w', encoding='utf-8') as file:
                        file.write(content)
                elif file_path.endswith('.docx'):
                    doc = Document()
                    doc.add_paragraph(content)
                    doc.save(file_path)
                elif file_path.endswith('.rtf'):
                    pypandoc.convert_text(content, 'rtf', format='md', outputfile=file_path, encoding='utf-8')
                elif file_path.endswith('.pdf'):
                    printer = QPrinter(QPrinter.HighResolution)
                    printer.setOutputFormat(QPrinter.PdfFormat)
                    printer.setOutputFileName(file_path)
                    doc = QTextDocument()
                    doc.setPlainText(content)
                    doc.print_(printer)
                else:
                    self.editor.error_dialog.show_error_popup("Unsupported file format")
            except Exception as e:
                self.editor.error_dialog.show_error_popup(str(e))

        # Update editor's current file path
        self.editor.current_file_path = file_path

