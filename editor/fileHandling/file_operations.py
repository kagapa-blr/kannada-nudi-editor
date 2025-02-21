import os

import pypandoc
from PyQt5.QtGui import QFontMetrics
from PyQt5.QtWidgets import QFileDialog, QMessageBox
from docx import Document

from utils.asciitounicode import process_line


class FileOperation:
    def __init__(self, editor):
        self.editor = editor
#----------------------------------------------------------------FILE OPEN----------------------------------------------------------------
    def checkOverflow(self):
        """Detect if text overflows and emit a signal for a new page."""
        if self.isOverflowing():
            self.textOverflow.emit()

    def isOverflowing(self):
        """Helper function to check if text overflows the available space."""
        font_metrics = QFontMetrics(self.editor.font())
        line_height = font_metrics.lineSpacing()
        document_margin = self.editor.contentsMargins()
        padding = 20  # Padding from setStyleSheet

        usable_height = self.editor.height() - document_margin.top() - document_margin.bottom() - (2 * padding)
        max_lines = usable_height // line_height
        current_lines = self.editor.document().blockCount()

        return current_lines >= max_lines

    def getMaxLinesPerPage(self):
        """Calculate the maximum number of lines that can fit in a page dynamically."""
        font_metrics = QFontMetrics(self.editor.font())
        line_height = font_metrics.lineSpacing()
        document_margin = self.editor.contentsMargins()
        padding = 20  # Padding from setStyleSheet

        usable_height = self.editor.height() - document_margin.top() - document_margin.bottom() - (2 * padding)
        return usable_height // line_height

    def addContentToPages(self, content):
        """Splits content into pages dynamically based on overflow detection."""
        lines = content.split("\n")
        current_page_content = []
        line_limit = self.getMaxLinesPerPage()

        for line in lines:
            current_page_content.append(line)
            if len(current_page_content) >= line_limit:
                self.editor.addPageWithContent("\n".join(current_page_content))
                current_page_content = []

        if current_page_content:
            self.editor.addPageWithContent("\n".join(current_page_content))

    def handle_open_file(self):
        options = QFileDialog.Options()
        self.editor.filename, _ = QFileDialog.getOpenFileName(
            self.editor,
            "Open File",
            "",
            "All Files (*);;Text Files (*.txt);;Word Documents (*.docx);;Rich Text Format (*.rtf)",
            options=options
        )

        if self.editor.filename:
            self.editor.content = ""
            file_extension = self.editor.filename.split('.')[-1].lower()

            try:
                if file_extension == 'txt':
                    with open(self.editor.filename, 'r', encoding="utf-8") as file:
                        content = file.read()
                elif file_extension == 'docx':
                    doc = Document(self.editor.filename)
                    content = "\n".join([para.text for para in doc.paragraphs])
                elif file_extension == 'rtf':
                    content = pypandoc.convert_file(self.editor.filename, 'plain', format='rtf')
                else:
                    raise ValueError("Unsupported file format")

                self.addContentToPages(content)
            except Exception as e:
                self.editor.error_dialog.show_error_popup(f"Error opening file: {str(e)}")

            self.editor.setWindowTitle("\u0c95\u0ca8\u0ccd\u0ca8\u0ca1 \u0ca8\u0cc1\u0ca1\u0cbf - " + self.editor.access_filename())
            self.editor.removeBlankPages()
            self.editor.statusbar.setStatusTip(f"Total pages: {self.editor.total_pages}")
            self.editor.current_file_path = self.editor.filename
    def handle_open_ascii_file(self):
        options = QFileDialog.Options()
        self.editor.filename, _ = QFileDialog.getOpenFileName(
            self.editor, "Open File", "",
            "All Files (*);;Text Files (*.txt);;Word Documents (*.docx);;Rich Text Format (*.rtf)",
            options=options)
        if self.editor.filename:
            content = ""
            file_extension = self.editor.filename.split('.')[-1].lower()

            try:
                if file_extension == 'txt':
                    with open(self.editor.filename, 'r', encoding="utf-8") as file:
                        unicode_lines = [process_line(line) for line in file]
                        content = ''.join(unicode_lines)
                elif file_extension == 'docx':
                    doc = Document(self.editor.filename)
                    paragraphs = [process_line(para.text) for para in doc.paragraphs]
                    content = "\n".join(paragraphs)
                elif file_extension == 'rtf':
                    plain_text = pypandoc.convert_file(self.editor.filename, 'plain', format='rtf')
                    unicode_lines = [process_line(line) for line in plain_text.splitlines()]
                    content = '\n'.join(unicode_lines)
                else:
                    self.editor.error_dialog("Unsupported file format")
                    # raise ValueError("Unsupported file format")

                # Split content into chunks of approximately 490 words
                words = content.split()
                word_limit = 490
                current_word_count = 0
                current_page_content = []
                for word in words:
                    current_page_content.append(word)
                    current_word_count += 1
                    if current_word_count >= word_limit:
                        self.editor.addPageWithContent(' '.join(current_page_content))
                        current_page_content = []
                        current_word_count = 0

                # Add any remaining content to a new page if not empty
                if current_page_content:
                    self.editor.addPageWithContent(' '.join(current_page_content))

            except Exception as e:
                self.editor.error_dialog.showError(str(e))

            # Update window title and remove blank pages
            self.editor.setWindowTitle("ಕನ್ನಡ ನುಡಿ - " + self.editor.access_filename())
            self.editor.removeBlankPages()

    # ----------------------------------------------------------------FILE OPEN----------------------------------------------------------------

    #----------------------------------------------------------------FILE SAVE----------------------------------------------------------------

    def handle_save_file(self):
        """Save the file while preserving formatting."""
        if self.editor.current_file_path:
            self.save_to_file(self.editor.current_file_path)
        else:
            self.handle_save_as_file()

    def handle_save_as_file(self):
        """Prompt user to select a file path and save the content with formatting."""
        options = QFileDialog.Options()
        file_path, selected_filter = QFileDialog.getSaveFileName(
            self.editor,
            "Save File As",
            "",
            "Word Documents (*.docx);;Rich Text Files (*.rtf *.html);;Plain Text Files (*.txt);;All Files (*)",
            options=options
        )

        if file_path:
            # Ensure correct file extension based on filter selection
            if "Word Documents" in selected_filter and not file_path.endswith(".docx"):
                file_path += ".docx"
            elif "Rich Text Files" in selected_filter and not (
                    file_path.endswith(".rtf") or file_path.endswith(".html")):
                file_path += ".rtf"
            elif "Plain Text Files" in selected_filter and not file_path.endswith(".txt"):
                file_path += ".txt"

            self.editor.current_file_path = file_path
            self.save_to_file(file_path)
            self.editor.setWindowTitle(f"ಕನ್ನಡ ನುಡಿ - {os.path.basename(file_path)}")

    def save_to_file(self, file_path):
        """Write content to file, preserving formatting."""
        try:
            if file_path.endswith(".docx"):
                self.write_docx(file_path)
            elif file_path.endswith((".rtf", ".html")):
                content = "\n\n".join([page.editor.toHtml() for page in self.editor.pages])
                with open(file_path, "w", encoding="utf-8") as file:
                    file.write(content)
            else:
                content = "\n\n".join([page.editor.toPlainText() for page in self.editor.pages])
                with open(file_path, "w", encoding="utf-8") as file:
                    file.write(content)

            QMessageBox.information(self.editor, "File Saved", "File saved successfully.")

        except Exception as e:
            QMessageBox.critical(self.editor, "Error", f"Failed to save file: {str(e)}")

    def read_docx(self, file_path):
        """Read a .docx file and convert it to HTML for display in the editor."""
        doc = Document(file_path)
        html_content = ""

        for para in doc.paragraphs:
            style = ""
            if para.runs:
                run = para.runs[0]
                if run.bold:
                    style += "font-weight:bold;"
                if run.italic:
                    style += "font-style:italic;"
                if run.underline:
                    style += "text-decoration:underline;"

            html_content += f"<p style='{style}'>{para.text}</p>"

        return html_content

    def write_docx(self, file_path):
        """Save editor content as a .docx file."""
        doc = Document()

        for page in self.editor.pages:
            text = page.editor.toPlainText()  # Extract plain text from the editor page
            paragraphs = text.splitlines()  # Use splitlines() to avoid extra blank lines

            for para in paragraphs:
                if para.strip():  # Avoid adding empty paragraphs
                    doc.add_paragraph(para)

        doc.save(file_path)  # Save the .docx file

#----------------------------------------------------------------FILE SAVE----------------------------------------------------------------
