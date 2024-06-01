import os
import subprocess

import docx
from PyQt5 import QtPrintSupport
from PyQt5 import QtWidgets, QtGui
from PyQt5.QtCore import QFileInfo, pyqtSlot
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QTextDocument, QTextDocumentWriter, QPainter, QTextCursor, QTextCharFormat, QIcon, \
    QTextBlockFormat
from PyQt5.QtPrintSupport import QPrinter
from PyQt5.QtWidgets import QFileDialog, QMenu, QMessageBox, QInputDialog, QLineEdit, QScrollArea
from PyQt5.QtWidgets import QWidget, QVBoxLayout, QLabel, QSlider, QHBoxLayout
from docx import Document

from config import file_path as fp
from editor.common_Dialogs import CommonDialogs
from editor.components.ascii_unicode_ConversionDialog import ConversionDialog
from editor.components.customise_page import Page
from editor.components.customize_image import ImageEditDialog
from editor.components.excel_csv_file_handling import ExcelCsvViewer
from editor.components.format_content import SpacingDialog
from editor.components.speech_to_text import SpeechToTextThread
from editor.components.table_functionality import TableFunctionality
from logger import setup_logger
from spellcheck.bloom_filter import bloom_lookup, reload_bloom_filter, start_bloom
from spellcheck.symspell_suggestions import suggestionReturner
from utils import find, datetime, table, wordcount
from utils.asciitounicode import process_line
from utils.corpus_clean import get_clean_words_for_dictionary
from utils.sort_by import SortDialog

filename = os.path.splitext(os.path.basename(__file__))[0]

# Set up logger
logger = setup_logger(filename)


def start_background_exe():
    exe_path = r"resources\keyboardDriver\kannadaKeyboard.exe"  # Path to your executable relative to the current directory
    #exe_path = r"resources\keyboardDriver\testing.exe"
    try:
        # Use subprocess.Popen to start the executable in the background
        print("kannada Nudi Keyboard loaded and running in background")
        subprocess.Popen([exe_path], shell=True, stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                         start_new_session=True)

    except Exception as e:
        print(f"Error starting background exe: {e}")



class TextEditor(QtWidgets.QMainWindow):

    def __init__(self, parent=None):
        super().__init__()
        QtWidgets.QMainWindow.__init__(self, parent)

        self.error_dialog = CommonDialogs()

        self.speech_thread = None
        self.scrollArea = None
        self.editor = None
        self.statusbar = None
        self.pages = []
        self.filename = None

        self.changesSaved = True

        self.initUI()




    def initToolbar(self):

        self.newAction = QtWidgets.QAction(QtGui.QIcon("resources/images/new-file.png"), "ಹೊಸ ಕಡತ", self)
        self.newAction.setShortcut("Ctrl+N")
        self.newAction.setStatusTip("Create a new document from scratch.")
        self.newAction.triggered.connect(self.new)

        self.openAction = QtWidgets.QAction(QtGui.QIcon("resources/images/open-file.png"), "ಕಡತ ತೆರೆಯಿರಿ", self)
        self.openAction.setStatusTip("Open existing document")
        self.openAction.setShortcut("Ctrl+O")
        self.openAction.triggered.connect(self.open)

        self.saveAction = QtWidgets.QAction(QtGui.QIcon("resources/images/stock_save.png"), "Save", self)
        self.saveAction.setStatusTip("Save document")
        self.saveAction.setShortcut("Ctrl+S")
        self.saveAction.triggered.connect(self.save)

        self.saveAsAction = QtWidgets.QAction(QtGui.QIcon("resources/images/stock_save.png"), "Save As...", self)
        self.saveAsAction.setStatusTip("Save document with a new name")
        self.saveAsAction.setShortcut("Ctrl+Shift+S")
        self.saveAsAction.triggered.connect(self.save_as)

        self.printAction = QtWidgets.QAction(QtGui.QIcon("resources/images/print.png"), "Print document", self)
        self.printAction.setStatusTip("Print document")
        self.printAction.setShortcut("Ctrl+P")
        self.printAction.triggered.connect(self.printHandler)

        self.previewAction = QtWidgets.QAction(QtGui.QIcon("resources/images/preview.png"), "Page view", self)
        self.previewAction.setStatusTip("Preview page before printing")
        self.previewAction.setShortcut("Ctrl+Shift+P")
        self.previewAction.triggered.connect(self.preview)

        self.findAction = QtWidgets.QAction(QtGui.QIcon("resources/images/find.png"), "Find and replace", self)
        self.findAction.setStatusTip("Find and replace words in your document")
        self.findAction.setShortcut("Ctrl+F")
        self.findAction.triggered.connect(find.Find(self).show)

        self.cutAction = QtWidgets.QAction(QtGui.QIcon("resources/images/stock_cut.png"), "Cut to clipboard", self)
        self.cutAction.setStatusTip("Delete and copy text to clipboard")
        self.cutAction.setShortcut("Ctrl+X")
        self.cutAction.triggered.connect(self.editor.cut)

        self.copyAction = QtWidgets.QAction(QtGui.QIcon("resources/images/stock_copy.png"), "Copy to clipboard", self)
        self.copyAction.setStatusTip("Copy text to clipboard")
        self.copyAction.setShortcut("Ctrl+C")
        self.copyAction.triggered.connect(self.editor.copy)

        self.pasteAction = QtWidgets.QAction(QtGui.QIcon("resources/images/stock_paste.png"), "Paste from clipboard",
                                             self)
        self.pasteAction.setStatusTip("Paste text from clipboard")
        self.pasteAction.setShortcut("Ctrl+V")
        self.pasteAction.triggered.connect(self.editor.paste)

        self.undoAction = QtWidgets.QAction(QtGui.QIcon("resources/images/undo.png"), "Undo last action", self)
        self.undoAction.setStatusTip("Undo last action")
        self.undoAction.setShortcut("Ctrl+Z")
        self.undoAction.triggered.connect(self.editor.undo)

        self.redoAction = QtWidgets.QAction(QtGui.QIcon("resources/images/redo.png"), "Redo last undone thing", self)
        self.redoAction.setStatusTip("Redo last undone thing")
        self.redoAction.setShortcut("Ctrl+Y")
        self.redoAction.triggered.connect(self.editor.redo)

        dateTimeAction = QtWidgets.QAction(QtGui.QIcon("resources/images/calender.png"), "Insert current date/time",
                                           self)
        dateTimeAction.setStatusTip("Insert current date/time")
        dateTimeAction.setShortcut("Ctrl+D")
        dateTimeAction.triggered.connect(datetime.DateTime(self).show)

        wordCountAction = QtWidgets.QAction(QtGui.QIcon("resources/images/count.png"), "See word/symbol count", self)
        wordCountAction.setStatusTip("See word/symbol count")
        wordCountAction.setShortcut("Ctrl+W")
        wordCountAction.triggered.connect(self.wordCount)

        tableAction = QtWidgets.QAction(QtGui.QIcon("resources/images/insert-table.png"), "Insert table", self)
        tableAction.setStatusTip("Insert table")
        tableAction.setShortcut("Ctrl+T")
        tableAction.triggered.connect(table.Table(self).show)

        imageAction = QtWidgets.QAction(QtGui.QIcon("resources/images/add-image.png"), "Insert image", self)
        imageAction.setStatusTip("Insert image")
        imageAction.setShortcut("Ctrl+Shift+I")
        imageAction.triggered.connect(self.choose_image)

        bulletAction = QtWidgets.QAction(QtGui.QIcon("resources/images/bullet-list.png"), "Insert bullet List", self)
        bulletAction.setStatusTip("Insert bullet list")
        bulletAction.setShortcut("Ctrl+Shift+B")
        bulletAction.triggered.connect(self.bulletList)

        numberedAction = QtWidgets.QAction(QtGui.QIcon("resources/images/number-list.png"), "Insert numbered List",
                                           self)
        numberedAction.setStatusTip("Insert numbered list")
        numberedAction.setShortcut("Ctrl+Shift+L")
        numberedAction.triggered.connect(self.numberList)

        sort_by_action = QtWidgets.QAction(QtGui.QIcon('resources/images/sortBy.png'), 'Sort By', self)
        sort_by_action.setStatusTip("Sort By")
        sort_by_action.triggered.connect(self.sort_by_action)

        speech_to_text = QtWidgets.QAction(QtGui.QIcon('resources/images/mic-speecch-to-text.png'), 'speech to Text',
                                           self)
        speech_to_text.setStatusTip("speech to text")
        speech_to_text.setCheckable(True)
        speech_to_text.triggered.connect(self.toggle_speech_to_text)

        ascii_to_unicode = QtWidgets.QAction(QtGui.QIcon('resources/images/convert.png'), 'speech to Text',
                                             self)
        ascii_to_unicode.setStatusTip("ASCII to Unicode vs converter")
        ascii_to_unicode.setCheckable(True)
        ascii_to_unicode.triggered.connect(self.ascii_to_unicode_converter)

        excel_csv = QtWidgets.QAction(QtGui.QIcon('resources/images/excel_csv.png'), 'Excel and CSV file operations',
                                      self)
        excel_csv.setStatusTip("ASCII to Unicode vs converter")
        excel_csv.setCheckable(True)
        excel_csv.triggered.connect(self.excel_csv_file)

        refresh_action = QtWidgets.QAction(QtGui.QIcon('resources/images/refresh.png'), 'Refresh and Recheck', self)
        refresh_action.setStatusTip("Refresh and Recheck")
        refresh_action.triggered.connect(self.refresh_recheck)

        self.toolbar = self.addToolBar("Options")
        self.toolbar.addAction(self.newAction)
        self.toolbar.addAction(self.openAction)
        self.toolbar.addAction(self.saveAction)
        self.toolbar.addAction(self.saveAsAction)
        self.toolbar.addSeparator()
        self.toolbar.addAction(self.printAction)
        self.toolbar.addAction(self.previewAction)
        self.toolbar.addSeparator()
        self.toolbar.addAction(self.cutAction)
        self.toolbar.addAction(self.copyAction)
        self.toolbar.addAction(self.pasteAction)
        self.toolbar.addAction(self.undoAction)
        self.toolbar.addAction(self.redoAction)
        self.toolbar.addSeparator()
        self.toolbar.addAction(self.findAction)
        self.toolbar.addAction(dateTimeAction)
        self.toolbar.addAction(wordCountAction)
        self.toolbar.addAction(tableAction)
        self.toolbar.addAction(imageAction)
        self.toolbar.addSeparator()
        self.toolbar.addAction(bulletAction)
        self.toolbar.addAction(numberedAction)
        self.toolbar.addAction(sort_by_action)
        self.toolbar.addAction(speech_to_text)
        self.toolbar.addAction(refresh_action)
        self.toolbar.addAction(ascii_to_unicode)
        self.toolbar.addAction(excel_csv)

        self.addToolBarBreak()

    def initFormatbar(self):

        fontBox = QtWidgets.QFontComboBox(self)
        fontBox.currentFontChanged.connect(lambda font: self.editor.setCurrentFont(font))

        # Set default font to "Nudist Parijath"
        default_font = QtGui.QFont("nudiParijatha")
        fontBox.setCurrentFont(default_font)

        fontSize = QtWidgets.QSpinBox(self)

        # Will display " pt" after each value
        fontSize.setSuffix(" pt")

        fontSize.valueChanged.connect(lambda size: self.editor.setFontPointSize(size))

        fontSize.setValue(12)

        fontColor = QtWidgets.QAction(QtGui.QIcon("resources/images/font-color.png"), "Change font color", self)
        fontColor.triggered.connect(self.fontColorChanged)

        boldAction = QtWidgets.QAction(QtGui.QIcon("resources/images/bold.png"), "Bold", self)
        boldAction.triggered.connect(self.bold)

        italicAction = QtWidgets.QAction(QtGui.QIcon("resources/images/italic.png"), "Italic", self)
        italicAction.triggered.connect(self.italic)

        underlAction = QtWidgets.QAction(QtGui.QIcon("resources/images/underline.png"), "Underline", self)
        underlAction.triggered.connect(self.underline)

        strikeAction = QtWidgets.QAction(QtGui.QIcon("resources/images/strike.png"), "Strike-out", self)
        strikeAction.triggered.connect(self.strike)

        superAction = QtWidgets.QAction(QtGui.QIcon("resources/images/superscript.png"), "Superscript", self)
        superAction.triggered.connect(self.superScript)

        subAction = QtWidgets.QAction(QtGui.QIcon("resources/images/subscript.png"), "Subscript", self)
        subAction.triggered.connect(self.subScript)

        alignLeft = QtWidgets.QAction(QtGui.QIcon("resources/images/align-left.png"), "Align left", self)
        alignLeft.triggered.connect(self.alignLeft)

        alignCenter = QtWidgets.QAction(QtGui.QIcon("resources/images/align-center.png"), "Align center", self)
        alignCenter.triggered.connect(self.alignCenter)

        alignRight = QtWidgets.QAction(QtGui.QIcon("resources/images/align-right.png"), "Align right", self)
        alignRight.triggered.connect(self.alignRight)

        alignJustify = QtWidgets.QAction(QtGui.QIcon("resources/images/align-justify.png"), "Align justify", self)
        alignJustify.triggered.connect(self.alignJustify)

        indentAction = QtWidgets.QAction(QtGui.QIcon("resources/images/indent.png"), "Indent Area", self)
        indentAction.setShortcut("Ctrl+Tab")
        indentAction.triggered.connect(self.indent)

        dedentAction = QtWidgets.QAction(QtGui.QIcon("resources/images/dedent.png"), "Dedent Area", self)
        dedentAction.setShortcut("Shift+Tab")
        dedentAction.triggered.connect(self.dedent)

        backColor = QtWidgets.QAction(QtGui.QIcon("resources/images/highlight.png"), "Change background color", self)
        backColor.triggered.connect(self.highlight)

        line_para_spacing = QtWidgets.QAction(QtGui.QIcon("resources/images/line-paragraph-spacing.png"),
                                              "line and Paragaraph spacing", self)
        line_para_spacing.setStatusTip("line and paragraph spacing .")
        line_para_spacing.triggered.connect(self.setSpacing)

        self.formatbar = self.addToolBar("Format")

        self.formatbar.addWidget(fontBox)
        self.formatbar.addWidget(fontSize)

        self.formatbar.addSeparator()

        self.formatbar.addAction(fontColor)
        self.formatbar.addAction(backColor)

        self.formatbar.addSeparator()

        self.formatbar.addAction(boldAction)
        self.formatbar.addAction(italicAction)
        self.formatbar.addAction(underlAction)
        self.formatbar.addAction(strikeAction)
        self.formatbar.addAction(superAction)
        self.formatbar.addAction(subAction)

        self.formatbar.addSeparator()

        self.formatbar.addAction(alignLeft)
        self.formatbar.addAction(alignCenter)
        self.formatbar.addAction(alignRight)
        self.formatbar.addAction(alignJustify)

        self.formatbar.addSeparator()

        self.formatbar.addAction(indentAction)
        self.formatbar.addAction(dedentAction)
        self.formatbar.addAction(line_para_spacing)

    def initMenubar(self):

        menubar = self.menuBar()

        file = menubar.addMenu("File")
        edit = menubar.addMenu("Edit")
        view = menubar.addMenu("View")

        # Add the most important actions to the menubar

        file.addAction(self.newAction)
        file.addAction(self.openAction)
        file.addAction(self.saveAction)
        file.addAction(self.saveAsAction)
        file.addAction(self.printAction)
        file.addAction(self.previewAction)

        edit.addAction(self.undoAction)
        edit.addAction(self.redoAction)
        edit.addAction(self.cutAction)
        edit.addAction(self.copyAction)
        edit.addAction(self.pasteAction)
        edit.addAction(self.findAction)

        # Toggling actions for the various bars
        toolbarAction = QtWidgets.QAction("Toggle Toolbar", self)
        toolbarAction.triggered.connect(self.toggleToolbar)

        formatbarAction = QtWidgets.QAction("Toggle Formatbar", self)
        formatbarAction.triggered.connect(self.toggleFormatbar)

        statusbarAction = QtWidgets.QAction("Toggle Statusbar", self)
        statusbarAction.triggered.connect(self.toggleStatusbar)

        view.addAction(toolbarAction)
        view.addAction(formatbarAction)
        view.addAction(statusbarAction)

    def initZoomSlider(self):
        # Create a zoom slider
        zoomLabel = QLabel("Zoom:", self)
        self.zoomSlider = QSlider(Qt.Horizontal, self)
        self.zoomSlider.setRange(10, 300)  # Zoom range from 10% to 300%
        self.zoomSlider.setValue(100)  # Default zoom level at 100%
        self.zoomSlider.setTickInterval(10)
        self.zoomSlider.setTickPosition(QSlider.TicksBelow)
        self.zoomSlider.setFixedWidth(150)  # Set a fixed width for the slider
        self.zoomSlider.valueChanged.connect(self.zoomText)

        # Add zoom slider to the status bar
        zoomWidget = QWidget(self)
        zoomLayout = QHBoxLayout(zoomWidget)
        zoomLayout.setContentsMargins(0, 0, 0, 0)  # No margins
        zoomLayout.addWidget(zoomLabel)
        zoomLayout.addWidget(self.zoomSlider)
        zoomWidget.setLayout(zoomLayout)

        self.statusbar.addPermanentWidget(zoomWidget)

    def initUI(self):
        #start_background_exe()
        self.scrollArea = QScrollArea(self)
        self.scrollArea.setWidgetResizable(True)
        centralWidget = QWidget()
        self.scrollArea.setWidget(centralWidget)
        self.setCentralWidget(self.scrollArea)

        # Create a QVBoxLayout for centering the pages
        centralLayout = QVBoxLayout(centralWidget)
        centralLayout.setAlignment(Qt.AlignCenter)  # Align the pages to the center

        self.pageLayout = QVBoxLayout()  # Remove the parent widget from the QVBoxLayout

        centralLayout.addLayout(self.pageLayout)  # Add the page layout to the central layout

        self.addPage()

        self.initToolbar()
        self.initFormatbar()
        self.initMenubar()

        # Initialize a statusbar for the window
        self.statusbar = self.statusBar()

        # Initialize the zoom slider
        self.initZoomSlider()

        self.setGeometry(100, 100, 1030, 800)
        self.setWindowTitle("ಕನ್ನಡ ನುಡಿ - " + self.access_filename())
        self.setWindowIcon(QIcon('resources/images/logo.jpg'))  # Set the application icon

        # Open in maximized mode
        self.showMaximized()

    def addPage(self):
        new_page = Page(self)
        self.pages.append(new_page)
        self.pageLayout.addWidget(new_page)
        self.pageLayout.addStretch()  # Add stretch to keep the pages at the top

        # Connect the signal to update the active editor
        new_page.activeEditorChanged.connect(self.updateActiveEditor)

        # Set the initial editor if this is the first page
        if len(self.pages) == 1:
            self.editor = new_page.editor

        self.editor.textChanged.connect(self.changed)
        self.editor.cursorPositionChanged.connect(self.cursorPosition)
        # We need our own context menu for tables
        self.editor.setContextMenuPolicy(Qt.CustomContextMenu)
        self.editor.customContextMenuRequested.connect(self.context)

        self.editor.installEventFilter(self)
        self.editor.setTabStopWidth(33)

        # Connect the text changed signal to check for overflow
        new_page.editor.textChanged.connect(lambda: self.checkPageOverflow(new_page))

    def zoomText(self, value):
        factor = value / 100.0
        for page in self.pages:
            page.setZoomFactor(factor)

    def updateActiveEditor(self, editor):
        self.editor = editor

    @pyqtSlot()
    def checkPageOverflow(self, page):
        editor = page.editor
        if editor.document().size().height() > editor.height():
            editor.textChanged.disconnect()
            self.addPage()

    def changed(self):
        self.changesSaved = False

    def closeEvent(self, event):
        try:
            if self.speech_thread:
                self.speech_thread.stop()
            event.accept()
        except Exception as e:
            self.error_dialog.show_error_popup(str(e))


        if self.changesSaved:
            event.accept()
        else:
            popup = QtWidgets.QMessageBox(self)
            popup.setIcon(QtWidgets.QMessageBox.Warning)
            popup.setWindowTitle("ಕನ್ನಡ ನುಡಿ")  # Set the title here
            popup.setText("ಡಾಕ್ಯುಮೆಂಟ್ ಅನ್ನು ಮಾರ್ಪಡಿಸಲಾಗಿದೆ")

            popup.setInformativeText("ನಿಮ್ಮ ಬದಲಾವಣೆಗಳನ್ನು ಉಳಿಸಲು ನೀವು ಬಯಸುವಿರಾ?")

            popup.setStandardButtons(QtWidgets.QMessageBox.Save |
                                     QtWidgets.QMessageBox.Cancel |
                                     QtWidgets.QMessageBox.Discard)

            popup.setDefaultButton(QtWidgets.QMessageBox.Save)

            answer = popup.exec_()

            if answer == QtWidgets.QMessageBox.Save:
                self.save()

            elif answer == QtWidgets.QMessageBox.Discard:
                event.accept()

            else:
                event.ignore()

    def context(self, pos):
        # Obtain the cursor at the position
        cursor = self.editor.cursorForPosition(pos)

        # Grab the current table, if there is one
        table = cursor.currentTable()

        # Above will return None if there is no current table
        if table:
            menu = QMenu(self)

            # Instantiate TableFunctionality for table manipulation
            table_functionality = TableFunctionality()

            # Add actions for table manipulation
            appendRowAction = QtWidgets.QAction("Append row", self)
            appendRowAction.triggered.connect(lambda: table.appendRows(1))
            appendColAction = QtWidgets.QAction("Append column", self)
            appendColAction.triggered.connect(lambda: table.appendColumns(1))
            removeRowAction = QtWidgets.QAction("Remove row", self)
            removeRowAction.triggered.connect(lambda: table_functionality.removeRow(table, cursor))
            removeColAction = QtWidgets.QAction("Remove column", self)
            removeColAction.triggered.connect(lambda: table_functionality.removeCol(table, cursor))
            insertRowAction = QtWidgets.QAction("Insert row", self)
            insertRowAction.triggered.connect(lambda: table_functionality.insertRow(table, cursor))
            insertColAction = QtWidgets.QAction("Insert column", self)
            insertColAction.triggered.connect(lambda: table_functionality.insertCol(table, cursor))
            mergeAction = QtWidgets.QAction("Merge cells", self)
            mergeAction.triggered.connect(lambda: table_functionality.mergeCells(table, cursor))
            splitAction = QtWidgets.QAction("Split cells", self)
            splitAction.triggered.connect(lambda: table_functionality.splitCell(table, cursor))

            # Only allow merging if there is a selection
            if not cursor.hasSelection():
                mergeAction.setEnabled(False)

            # Add actions to the menu
            menu.addAction(appendRowAction)
            menu.addAction(appendColAction)
            menu.addSeparator()
            menu.addAction(removeRowAction)
            menu.addAction(removeColAction)
            menu.addSeparator()
            menu.addAction(insertRowAction)
            menu.addAction(insertColAction)
            menu.addSeparator()
            menu.addAction(mergeAction)
            menu.addAction(splitAction)

            # Show the menu
            menu.exec_(self.editor.mapToGlobal(pos))

        else:
            # Handle word suggestions if no table is found
            cursor.select(QTextCursor.WordUnderCursor)
            selected_word = cursor.selectedText()
            if selected_word:
                menu = QMenu(self)
                if bloom_lookup(selected_word):
                    return
                suggestions = suggestionReturner(selected_word)

                # Create a function to handle word replacement
                def replace_word_with_suggestion(suggestion):
                    cursor.removeSelectedText()
                    cursor.insertText(suggestion)

                # Add each suggestion to the context menu and connect them to the replacement function
                for suggestion in suggestions:
                    action = menu.addAction(suggestion)
                    action.triggered.connect(lambda _, s=suggestion: replace_word_with_suggestion(s))

                # Add actions for adding to dictionary, ignoring, and replacing
                add_to_dict_action = menu.addAction("Add to Dictionary")
                ignore_action = menu.addAction("Ignore")
                replace_action = menu.addAction("Replace All")

                # Connect actions to corresponding functions
                add_to_dict_action.triggered.connect(lambda: self.addToDictionary(selected_word))
                ignore_action.triggered.connect(lambda: self.ignore_word())
                replace_action.triggered.connect(lambda: self.replace_word(selected_word))

                # Show the menu
                menu.exec_(self.editor.mapToGlobal(pos))
            else:
                # If no selection, call the default context menu event
                event = QtGui.QContextMenuEvent(QtGui.QContextMenuEvent.Mouse, pos)
                self.editor.contextMenuEvent(event)



    def toggleToolbar(self):

        state = self.toolbar.isVisible()

        # Set the visibility to its inverse
        self.toolbar.setVisible(not state)

    def toggleFormatbar(self):

        state = self.formatbar.isVisible()

        # Set the visibility to its inverse
        self.formatbar.setVisible(not state)

    def toggleStatusbar(self):

        state = self.statusbar.isVisible()

        # Set the visibility to its inverse
        self.statusbar.setVisible(not state)

    def new(self):

        spawn = TextEditor()

        spawn.show()

    def open(self):
        # Define the file filter
        filter = "Text Files (*.txt);;Word Documents (*.docx);;Rich Text Files (*.rtf)"
        # Get filename using QFileDialog
        self.filename, _ = QFileDialog.getOpenFileName(self, 'Open File', ".", filter)
        print("filename: ", self.filename)

        if self.filename:
            # Open and read the selected file based on its extension
            try:
                if self.filename.endswith('.txt'):
                    with open(self.filename, 'r', encoding='utf-8') as file:
                        unicode_lines = [process_line(line) for line in file.readlines()]
                        self.editor.setText('\n'.join(unicode_lines))
                elif self.filename.endswith('.rtf'):
                    with open(self.filename, 'r') as file:
                        rtf_text = file.read()
                        document = QTextDocument()
                        document.setHtml(rtf_text)
                        self.editor.setDocument(document)
                elif self.filename.endswith('.docx'):
                    # Use python-docx to extract text from DOCX with formatting
                    doc = docx.Document(self.filename)
                    docx_text = []
                    for para in doc.paragraphs:
                        docx_text.append(para.text)
                    # Join the paragraphs into a single string with newlines
                    docx_text = "\n".join(docx_text)
                    self.editor.setText(docx_text)
                else:
                    QMessageBox.critical(self, 'Error', 'Unsupported file format')
            except Exception as e:
                QMessageBox.critical(self, 'Error', f'Error opening file: {str(e)}')

        self.setWindowTitle("ಕನ್ನಡ ನುಡಿ - " + self.access_filename())

    def save(self):
        try:
            if not self.filename:
                # If no filename is set, prompt the user to choose a file location for saving
                self.filename, _ = QFileDialog.getSaveFileName(self, 'Save File', '',
                                                               filter="All Files (*);;RTF Files (*.rtf);;DOCX Files (*.docx);;PDF Files (*.pdf);;Text Files (*.txt)")

                if not self.filename:  # User cancelled the dialog
                    return

                # Determine the file extension and set the appropriate filter
                file_info = QFileInfo(self.filename)
                file_suffix = file_info.suffix().lower()
                if file_suffix not in ['rtf', 'docx', 'pdf', 'txt']:
                    QMessageBox.critical(self, 'Error', 'Unsupported file format.')
                    return

            # Create a QTextDocument and set the HTML content from QTextEdit
            document = QTextDocument()
            document.setHtml(self.editor.toHtml())

            if self.filename.endswith(".rtf"):
                # Save as RTF using QTextDocumentWriter
                writer = QTextDocumentWriter(self.filename)
                success = writer.write(document)
            elif self.filename.endswith(".docx"):
                # Save as DOCX using python-docx
                document = Document()
                for line in self.editor.toPlainText().split('\n'):
                    document.add_paragraph(line)
                document.save(self.filename)  # Use self.filename instead of filename
                success = True
            elif self.filename.endswith(".pdf"):
                # Save as PDF using QPrinter and QPainter
                printer = QPrinter(QPrinter.HighResolution)
                printer.setOutputFormat(QPrinter.PdfFormat)
                printer.setOutputFileName(self.filename)
                painter = QPainter(printer)
                document.drawContents(painter)
                painter.end()
                success = True
            elif self.filename.endswith(".txt"):
                # Save as TXT
                with open(self.filename, "w", encoding="utf-8") as txt_file:
                    txt_file.write(self.editor.toPlainText())
                success = True
            else:
                QMessageBox.critical(self, 'Error', 'Unsupported file format.')
                return

            if success:
                QMessageBox.information(self, 'Success', f'File saved successfully: {self.filename}')
            else:
                QMessageBox.critical(self, 'Error', f'Error saving file: {self.filename}')

        except Exception as e:
            print(str(e))
            QMessageBox.critical(self, 'Error', f'Error saving file: {str(e)}')

        self.setWindowTitle("ಕನ್ನಡ ನುಡಿ - " + self.access_filename())

    def save_as(self):
        try:
            # Prompt the user to choose a file location for saving
            filename, _ = QFileDialog.getSaveFileName(self, 'Save File As', '',
                                                      filter="All Files (*);;RTF Files (*.rtf);;DOCX Files (*.docx);;PDF Files (*.pdf);;Text Files (*.txt)")

            if not filename:  # User cancelled the dialog
                return

            # Determine the file extension and set the appropriate filter
            file_info = QFileInfo(filename)
            file_suffix = file_info.suffix().lower()
            if file_suffix not in ['rtf', 'docx', 'pdf', 'txt']:
                QMessageBox.critical(self, 'Error', 'Unsupported file format.')
                return

            if filename.endswith(".rtf"):
                # Save as RTF using QTextDocumentWriter
                document = QTextDocument()
                document.setHtml(self.editor.toHtml())
                writer = QTextDocumentWriter(filename)
                success = writer.write(document)
            elif filename.endswith(".docx"):
                # Save as DOCX using python-docx
                document = Document()
                for line in self.editor.toPlainText().split('\n'):
                    document.add_paragraph(line)
                document.save(filename)
                success = True
            elif filename.endswith(".pdf"):
                # Save as PDF using QPrinter and QPainter
                printer = QPrinter(QPrinter.HighResolution)
                printer.setOutputFormat(QPrinter.PdfFormat)
                printer.setOutputFileName(filename)
                painter = QPainter(printer)
                document = QTextDocument()
                document.setHtml(self.editor.toHtml())
                document.drawContents(painter)
                painter.end()
                success = True
            elif filename.endswith(".txt"):
                # Save as TXT with UTF-8 encoding
                with open(filename, "w", encoding="utf-8") as txt_file:
                    txt_file.write(self.editor.toPlainText())
                success = True
            else:
                QMessageBox.critical(self, 'Error', 'Unsupported file format.')
                return

            if success:
                QMessageBox.information(self, 'Success', f'File saved successfully: {filename}')
            else:
                QMessageBox.critical(self, 'Error', f'Error saving file: {filename}')

        except Exception as e:
            print(str(e))
            QMessageBox.critical(self, 'Error', f'Error saving file: {str(e)}')

    def preview(self):

        # Open preview dialog
        preview = QtPrintSupport.QPrintPreviewDialog()

        # If a print is requested, open print dialog
        preview.paintRequested.connect(lambda p: self.editor.print_(p))

        preview.exec_()

    def printHandler(self):

        # Open printing dialog
        dialog = QtPrintSupport.QPrintDialog()

        if dialog.exec_() == QtWidgets.QDialog.Accepted:
            self.editor.document().print_(dialog.printer())

    def cursorPosition(self):

        cursor = self.editor.textCursor()

        # Mortals like 1-indexed things
        line = cursor.blockNumber() + 1
        col = cursor.columnNumber()

        self.statusbar.showMessage("Line: {} | Column: {}".format(line, col))

    def wordCount(self):

        wc = wordcount.WordCount(self)

        wc.getText()

        wc.show()

    def insertImage(self):

        # Get image file name
        # PYQT5 Returns a tuple in PyQt5
        filename = \
            QtWidgets.QFileDialog.getOpenFileName(self, 'Insert image', ".", "Images (*.png *.xpm *.jpg *.bmp *.gif)")[
                0]

        if filename:

            # Create image object
            image = QtGui.QImage(filename)

            # Error if unloadable
            if image.isNull():

                popup = QtWidgets.QMessageBox(QtWidgets.QMessageBox.Critical,
                                              "Image load error",
                                              "Could not load image file!",
                                              QtWidgets.QMessageBox.Ok,
                                              self)
                popup.show()

            else:

                cursor = self.editor.textCursor()

                cursor.insertImage(image, filename)

    def fontColorChanged(self):

        # Get a color from the text dialog
        color = QtWidgets.QColorDialog.getColor()

        # Set it as the new text color
        self.editor.setTextColor(color)

    def highlight(self):

        color = QtWidgets.QColorDialog.getColor()

        self.editor.setTextBackgroundColor(color)

    def bold(self):

        if self.editor.fontWeight() == QtGui.QFont.Bold:

            self.editor.setFontWeight(QtGui.QFont.Normal)

        else:

            self.editor.setFontWeight(QtGui.QFont.Bold)

    def italic(self):

        state = self.editor.fontItalic()

        self.editor.setFontItalic(not state)

    def underline(self):

        state = self.editor.fontUnderline()

        self.editor.setFontUnderline(not state)

    def strike(self):

        # Grab the text's format
        fmt = self.editor.currentCharFormat()

        # Set the fontStrikeOut property to its opposite
        fmt.setFontStrikeOut(not fmt.fontStrikeOut())

        # And set the next char format
        self.editor.setCurrentCharFormat(fmt)

    def superScript(self):

        # Grab the current format
        fmt = self.editor.currentCharFormat()

        # And get the vertical alignment property
        align = fmt.verticalAlignment()

        # Toggle the state
        if align == QtGui.QTextCharFormat.AlignNormal:

            fmt.setVerticalAlignment(QtGui.QTextCharFormat.AlignSuperScript)

        else:

            fmt.setVerticalAlignment(QtGui.QTextCharFormat.AlignNormal)

        # Set the new format
        self.editor.setCurrentCharFormat(fmt)

    def subScript(self):

        # Grab the current format
        fmt = self.editor.currentCharFormat()

        # And get the vertical alignment property
        align = fmt.verticalAlignment()

        # Toggle the state
        if align == QtGui.QTextCharFormat.AlignNormal:

            fmt.setVerticalAlignment(QtGui.QTextCharFormat.AlignSubScript)

        else:

            fmt.setVerticalAlignment(QtGui.QTextCharFormat.AlignNormal)

        # Set the new format
        self.editor.setCurrentCharFormat(fmt)

    def alignLeft(self):
        self.editor.setAlignment(Qt.AlignLeft)

    def alignRight(self):
        self.editor.setAlignment(Qt.AlignRight)

    def alignCenter(self):
        self.editor.setAlignment(Qt.AlignCenter)

    def alignJustify(self):
        self.editor.setAlignment(Qt.AlignJustify)

    def indent(self):

        # Grab the cursor
        cursor = self.editor.textCursor()

        if cursor.hasSelection():

            # Store the current line/block number
            temp = cursor.blockNumber()

            # Move to the selection's end
            cursor.setPosition(cursor.anchor())

            # Calculate range of selection
            diff = cursor.blockNumber() - temp

            direction = QtGui.QTextCursor.Up if diff > 0 else QtGui.QTextCursor.Down

            # Iterate over lines (diff absolute value)
            for n in range(abs(diff) + 1):
                # Move to start of each line
                cursor.movePosition(QtGui.QTextCursor.StartOfLine)

                # Insert tabbing
                cursor.insertText("\t")

                # And move back up
                cursor.movePosition(direction)

        # If there is no selection, just insert a tab
        else:

            cursor.insertText("\t")

    def handleDedent(self, cursor):

        cursor.movePosition(QtGui.QTextCursor.StartOfLine)

        # Grab the current line
        line = cursor.block().text()

        # If the line starts with a tab character, delete it
        if line.startswith("\t"):

            # Delete next character
            cursor.deleteChar()

        # Otherwise, delete all spaces until a non-space character is met
        else:
            for char in line[:8]:

                if char != " ":
                    break

                cursor.deleteChar()

    def dedent(self):

        cursor = self.editor.textCursor()

        if cursor.hasSelection():

            # Store the current line/block number
            temp = cursor.blockNumber()

            # Move to the selection's last line
            cursor.setPosition(cursor.anchor())

            # Calculate range of selection
            diff = cursor.blockNumber() - temp

            direction = QtGui.QTextCursor.Up if diff > 0 else QtGui.QTextCursor.Down

            # Iterate over lines
            for n in range(abs(diff) + 1):
                self.handleDedent(cursor)

                # Move up
                cursor.movePosition(direction)

        else:
            self.handleDedent(cursor)

    def bulletList(self):

        cursor = self.editor.textCursor()

        # Insert bulleted list
        cursor.insertList(QtGui.QTextListFormat.ListDisc)

    def numberList(self):

        cursor = self.editor.textCursor()

        # Insert list with numbers
        cursor.insertList(QtGui.QTextListFormat.ListDecimal)

    def addToDictionary(self, word):
        with open(fp.bloomfilter_data, 'a', encoding='utf-8') as dict_file:
            word = get_clean_words_for_dictionary(word)
            if not bloom_lookup(word):
                dict_file.write(word + '\n')
                with open(fp.symspell_word_freq_data, 'a', encoding='utf-8') as file:
                    file.write(f"{word} {1}\n")
                    print(word, "successfully Added to Dictionary")

                # Remove underline from the word
                cursor = self.editor.textCursor()
                cursor.beginEditBlock()  # Begin editing block to improve performance
                format = QTextCharFormat()
                format.setForeground(self.editor.palette().color(self.editor.foregroundRole()))
                format.setUnderlineStyle(QTextCharFormat.NoUnderline)  # Clear underline
                cursor.mergeCharFormat(format)
                cursor.endEditBlock()
            else:
                print("word already present in Dictionary file")

    def ignore_word(self):
        cursor = self.editor.textCursor()
        cursor.beginEditBlock()  # Begin editing block to improve performance

        format = QTextCharFormat()
        format.setForeground(self.editor.palette().color(self.editor.foregroundRole()))

        # Clear underline
        format.setUnderlineStyle(QTextCharFormat.NoUnderline)

        cursor.mergeCharFormat(format)
        cursor.endEditBlock()

    def replace_word(self, selected_word):
        cursor = self.editor.textCursor()
        if not cursor.hasSelection():
            return  # No text selected, nothing to replace

        new_word, ok = QInputDialog.getText(self, 'Replace Word', f'Enter the new word to replace "{selected_word}":',
                                            QLineEdit.Normal)
        if ok and new_word:
            cursor.insertText(new_word)
            cursor.removeSelectedText()

    def refresh_recheck(self):
        reload_bloom_filter()
        cursor_position = self.editor.textCursor()
        text_content = self.editor.toHtml()
        plain_text = self.editor.toPlainText()
        content_for_bloom = [get_clean_words_for_dictionary(word) for word in plain_text.split() if word if
                             len(word) > 1]
        wrong_words = start_bloom(content_for_bloom)
        # Highlight incorrect words in the editor
        highlighted_content = text_content
        for word in wrong_words:
            highlighted_content = highlighted_content.replace(word,
                                                              f'<span style="text-decoration: underline;">{word}</span>')
        # Update the editor with the highlighted content
        self.editor.setHtml(highlighted_content)
        # Restore the cursor position
        self.editor.setTextCursor(cursor_position)

    def access_filename(self):
        return self.filename if self.filename is not None else "Untitled"

    def choose_image(self):
        file_dialog = QFileDialog(self)
        file_dialog.setNameFilter("Images (*.png *.jpg *.bmp)")
        if file_dialog.exec_():
            file_path = file_dialog.selectedFiles()[0]
            self.editAndInsertImage(file_path)

    def editAndInsertImage(self, image_path):
        image_dialog = ImageEditDialog(image_path, self)
        if image_dialog.exec_():
            modified_image = image_dialog.getModifiedImage()
            if not modified_image.isNull():
                cursor = self.editor.textCursor()
                cursor.insertImage(modified_image)

    def sort_by_action(self):
        sort_dialog = SortDialog(self)
        if sort_dialog.exec_():
            sort_by = sort_dialog.combo_sort_by.currentText()
            type_ = sort_dialog.combo_type.currentText()
            using = sort_dialog.combo_using.currentText()
            ascending = sort_dialog.radio_asc.isChecked()
            has_header = sort_dialog.check_has_header.isChecked()
            separator = sort_dialog.line_separator.text()
            sort_options = sort_dialog.combo_sort_options.currentText()

            # Get text from editor
            text = self.editor.toPlainText()

            # Split text into lines (or paragraphs)
            lines = text.split('\n')

            # Implement your sorting logic here based on the user input
            # This example will sort the lines alphabetically
            if type_ == "Text":
                lines.sort(key=str.lower if not sort_options == "Case sensitive" else str, reverse=not ascending)
            elif type_ == "Number":
                lines.sort(key=lambda x: float(x) if x.replace('.', '', 1).isdigit() else float('inf'),
                           reverse=not ascending)
            elif type_ == "Date":
                from datetime import datetime
                lines.sort(key=lambda x: datetime.strptime(x, '%Y-%m-%d'), reverse=not ascending)

            sorted_text = '\n'.join(lines)

            # Set sorted text back to editor
            self.editor.setPlainText(sorted_text)

    def setSpacing(self):
        dialog = SpacingDialog(self)
        dialog.setWindowTitle('Line & Paragraph Spacing')
        dialog.lineSpacingLabel.setText('Line Spacing (pixels):')
        dialog.customLineSpacingSpinBox.setValue(self.editor.currentFont().pixelSize())
        dialog.beforeParagraphSpinBox.setValue(0)
        dialog.afterParagraphSpinBox.setValue(0)
        dialog.exec_()
        dialog.applySettings()

    def setLineSpacing(self, value):
        cursor = self.editor.textCursor()
        fmt = cursor.blockFormat()
        fmt.setLineHeight(value, QTextBlockFormat.LineDistanceHeight)
        cursor.setBlockFormat(fmt)

    def setParagraphSpacing(self, before, after):
        cursor = self.editor.textCursor()
        fmt = cursor.blockFormat()
        fmt.setTopMargin(before)
        fmt.setBottomMargin(after)
        cursor.setBlockFormat(fmt)

    def toggle_speech_to_text(self):
        sender = self.sender()
        if sender.isChecked():
            sender.setText("Stop Speech to Text")
            self.speech_thread = SpeechToTextThread(self.editor)
            self.speech_thread.start()
        else:
            sender.setText("Speech to Text")
            if self.speech_thread:
                self.speech_thread.stop()
                self.speech_thread = None

    def ascii_to_unicode_converter(self):
        dialog = ConversionDialog(self)
        dialog.exec_()

    def excel_csv_file(self):
        self.viewer = ExcelCsvViewer()
        self.viewer.show()
