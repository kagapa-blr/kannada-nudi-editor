import re
import sys

import docx
from PyQt5.QtCore import Qt, QEvent
from PyQt5.QtGui import QTextCursor, QTextCharFormat, QIcon, QFontDatabase
from PyQt5.QtWidgets import QApplication, QMainWindow, QTextEdit, QAction, QFileDialog, QMenu, QToolBar, QMessageBox, \
    QInputDialog, QLineEdit

import file_path as fp
from bloom_filter import bloom_lookup, start_bloom, reload_bloom_filter
from corpus_clean import get_clean_words_for_dictionary
from symspell_suggestions import suggestionReturner


class TextEditor(QMainWindow):
    def __init__(self):
        super().__init__()

        self.initUI()

    def initUI(self):
        self.text_edit = QTextEdit(self)
        self.setCentralWidget(self.text_edit)
        self.text_edit.setContextMenuPolicy(Qt.CustomContextMenu)
        self.text_edit.customContextMenuRequested.connect(self.showContextMenu)
        # Create a QFont with the "Noto Sans Kannada" font family
        font_url = "static/Nudi_fonts/NudiUni01e.ttf"
        self.font_id = QFontDatabase.addApplicationFont(font_url)
        self.family = QFontDatabase.applicationFontFamilies(self.font_id)[0]
        self.text_edit.installEventFilter(self)

        current_font = self.text_edit.currentFont()
        new_font_size = 15
        new_font = current_font
        new_font.setPointSize(new_font_size)
        new_font.setFamily(self.family)
        self.text_edit.setFont(new_font)
        # ----------------------------------------------------------------file Menu Bar start----------------------------------------------------------------
        menubar = self.menuBar()

        # File menu
        file_menu = menubar.addMenu('File')

        open_action = QAction(QIcon('images/stock_open-16.png'), 'Open', self)
        open_action.triggered.connect(self.openFile)
        file_menu.addAction(open_action)

        save_action = QAction(QIcon('images/stock_save.png'),'Save', self)
        save_action.triggered.connect(self.saveFile)
        file_menu.addAction(save_action)

        save_as_action = QAction(QIcon('images/stock_save-16.png'),'Save As', self)
        save_as_action.triggered.connect(self.saveFileAs)
        file_menu.addAction(save_as_action)

        save_as_docx_action = QAction('Save As .docx', self)
        save_as_docx_action.triggered.connect(self.saveFileAsDocx)
        file_menu.addAction(save_as_docx_action)

        # Edit menu
        edit_menu = menubar.addMenu('Edit')
        cut_action = QAction(QIcon('images/stock_cut.png'),'Cut', self)
        cut_action.triggered.connect(self.text_edit.cut)
        edit_menu.addAction(cut_action)

        copy_action = QAction(QIcon('images/stock_copy.png'),'Copy', self)
        copy_action.triggered.connect(self.text_edit.copy)
        edit_menu.addAction(copy_action)

        paste_action = QAction(QIcon('images/stock_paste.png'),'Paste', self)
        paste_action.triggered.connect(self.text_edit.paste)
        edit_menu.addAction(paste_action)

        edit_undo_action =QAction(QIcon('images/undo.png'), 'Undo', self)
        edit_undo_action.triggered.connect(self.text_edit.undo)
        edit_menu.addAction(edit_undo_action)

        edit_redo_action =QAction(QIcon('images/redo.png'),'Redo', self)
        edit_redo_action.triggered.connect(self.text_edit.redo)
        edit_menu.addAction(edit_redo_action)

        # Font size menu
        font_size_menu = menubar.addMenu('Font')

        increase_font_size_action = QAction('Increase Font Size (+)', self)
        increase_font_size_action.setShortcut('Ctrl++')  # Set the shortcut (e.g., Ctrl++)
        increase_font_size_action.triggered.connect(self.increaseFontSize)
        font_size_menu.addAction(increase_font_size_action)

        decrease_font_size_action = QAction('Decrease Font Size (-)', self)
        decrease_font_size_action.setShortcut('Ctrl+-')  # Set the shortcut (e.g., Ctrl+-)
        decrease_font_size_action.triggered.connect(self.decreaseFontSize)
        font_size_menu.addAction(decrease_font_size_action)
        # Add more font actions as needed

#----------------------------------------------------------------tool bar----------------------------------------------------------------
        toolbar = QToolBar()
        self.addToolBar(toolbar)
        # Add actions with icons to the toolbar
        open_action = QAction(QIcon('images/stock_open.png'), 'Open', self)
        open_action.triggered.connect(self.openFile)
        toolbar.addAction(open_action)

        save_action = QAction(QIcon('images/stock_save.png'),'Save', self)
        save_action.triggered.connect(self.saveFile)
        toolbar.addAction(save_action)

        refresh_action = QAction(QIcon('images/refresh.png'), 'Refresh and Recheck', self)
        refresh_action.triggered.connect(self.refresh_recheck)
        toolbar.addAction(refresh_action)

        undo_action =QAction(QIcon('images/undo.png'), 'Undo', self)
        undo_action.triggered.connect(self.text_edit.undo)
        toolbar.addAction(undo_action)

        redo_action =QAction(QIcon('images/redo.png'),'Redo', self)
        redo_action.triggered.connect(self.text_edit.redo)
        toolbar.addAction(redo_action)

#----------------------------------------------------------------Editor Size Settings----------------------------------------------------------------
        self.setGeometry(100, 100, 800, 600)
        self.setWindowTitle('Kannada Spellchecker')
        self.setWindowIcon(QIcon('images/logo.jpg'))  # Set the application icon
        self.show()




    def openFile(self):
        format = QTextCharFormat()
        format.setForeground(Qt.red)  # You can set the color to any color you want
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getOpenFileName(self, 'Open File', '','All Files (*);;Word Documents (*.docx);Text Files (*.txt);', options=options)
        if file_name:
            if file_name.endswith('.docx'):
                doc = docx.Document(file_name)
                paragraphs = [paragraph.text for paragraph in doc.paragraphs]
                content = "<br>".join(paragraphs)  # Preserve line breaks
            else:
                with open(file_name, 'r', encoding='utf-8') as file:
                    content = file.read()
            import time
            start_time = time.time()
            content_for_bloom = [get_clean_words_for_dictionary(word) for word in content.split() if word if len(word) > 1]
            print("clean_up time: ", time.time() - start_time)
            start_time = time.time()
            wrong_words = start_bloom(content_for_bloom)
            print("time taken to find wrong words: ", time.time() - start_time)
            start_time = time.time()
            for word in wrong_words:
                content = content.replace(word, f'<span style="color:red">{word}</span>')
            print("time taken to hilight words : ",time.time()-start_time)
            self.setWindowTitle('Kannada Spellchecker - ' + file_name)
            current_font = self.text_edit.currentFont()
            new_font_size = 12
            new_font = current_font
            new_font.setPointSize(new_font_size)
            new_font.setFamily(self.family)
            self.text_edit.setFont(new_font)
            self.text_edit.setHtml(content)


    def saveFile(self):
        if hasattr(self, 'current_file'):
            with open(self.current_file, 'w', encoding='utf-8') as file:
                if(len(self.text_edit.toPlainText())==0):
                    print("No content")
                file.write(self.text_edit.toPlainText())
        else:
            self.saveFileAs()

    def saveFileAs(self):
        content = self.text_edit.toPlainText()
        if len(content) == 0:
            QMessageBox.warning(self, 'Warning', 'No content to save.')
            return
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(self, 'Save File', '', 'Text Files (*.txt);;All Files (*)', options=options)
        if file_name:
            with open(file_name, 'w', encoding='utf-8') as file:
                file.write(self.text_edit.toPlainText())
            self.current_file = file_name

    def saveFileAsDocx(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(self, 'Save As .docx', '', 'Word Documents (*.docx);;All Files (*)', options=options)
        if file_name:
            doc = docx.Document()
            doc.add_paragraph(self.text_edit.toPlainText())
            doc.save(file_name)

    def showContextMenu(self, pos):
        cursor = self.text_edit.cursorForPosition(pos)
        cursor.select(QTextCursor.WordUnderCursor)
        selected_word = cursor.selectedText()
        if  bloom_lookup(selected_word):
            return
        suggestions = suggestionReturner(selected_word)
        if selected_word:
            menu = QMenu(self)

            # Create a function to handle the word replacement
            def replace_word_with_suggestion(suggestion):
                cursor.removeSelectedText()
                cursor.insertText(suggestion)

            # Add each suggestion to the context menu and connect them to the replacement function
            for suggestion in suggestions:
                action = menu.addAction(suggestion)
                action.triggered.connect(lambda _, s=suggestion: replace_word_with_suggestion(s))

            add_to_dict_action = menu.addAction("Add to Dictionary")
            ignore_action = menu.addAction("Ignore")
            replace_action = menu.addAction("Replace All")

            # Connect actions to corresponding functions
            add_to_dict_action.triggered.connect(lambda: self.addToDictionary(selected_word))
            ignore_action.triggered.connect(lambda : self.ignore_word(selected_word))
            replace_action.triggered.connect(lambda : self.replace_word(selected_word))

            menu.exec_(self.text_edit.mapToGlobal(pos))

    def addToDictionary(self, word):
        with open(fp.bloomfilter_data, 'a', encoding='utf-8') as dict_file:

            word = get_clean_words_for_dictionary(word)
            if not bloom_lookup(word):
                dict_file.write(word + '\n')
                with open(fp.symspell_word_freq_data, 'a', encoding='utf-8') as file:
                        file.write(f"{word} {1}\n")
                        print(word,"successfully Added to Dictionary")
                cursor = self.text_edit.textCursor()
                cursor.beginEditBlock()  # Begin editing block to improve performance
                format = QTextCharFormat()
                format.setForeground(self.text_edit.palette().color(self.text_edit.foregroundRole()))
                cursor.mergeCharFormat(format)
                cursor.endEditBlock()
            else:
                print("word already present in Dictionary file")


    def ignore_word(self,word):
        cursor = self.text_edit.textCursor()
        cursor.beginEditBlock()  # Begin editing block to improve performance
        format = QTextCharFormat()
        format.setForeground(self.text_edit.palette().color(self.text_edit.foregroundRole()))
        cursor.mergeCharFormat(format)
        cursor.endEditBlock()


    def replace_word(self, word_to_replace):
        new_word, ok = QInputDialog.getText(self, 'Replace Word', f'Enter new word to replace {word_to_replace}',
                                            QLineEdit.Normal)
        if ok and new_word:
            cursor = self.text_edit.textCursor()  # Get the current cursor position
            current_text = self.text_edit.toPlainText()
            new_text = current_text.replace(word_to_replace, new_word)
            self.text_edit.setPlainText(new_text)
            # Restore the cursor position
            cursor.setPosition(cursor.position() + len(new_word))
            self.text_edit.setTextCursor(cursor)

    def increaseFontSize(self):
        current_font = self.text_edit.currentFont()
        new_font_size = current_font.pointSize() + 2
        new_font = current_font
        new_font.setPointSize(new_font_size)
        new_font.setFamily(self.family)
        self.text_edit.setFont(new_font)

    def decreaseFontSize(self):
        current_font = self.text_edit.currentFont()
        new_font_size = max(current_font.pointSize() - 2, 2)  # Ensure font size doesn't go below 2
        new_font = current_font
        new_font.setPointSize(new_font_size)
        new_font.setFamily(self.family)
        self.text_edit.setFont(new_font)
    def changeFontSize(self, delta):
        current_font = self.text_edit.currentFont()
        new_font_size = max(current_font.pointSize() + delta, 2)
        new_font = current_font
        new_font.setPointSize(new_font_size)
        self.text_edit.setFont(new_font)
    def enlargeText(self):
        self.changeFontSize(2)

    def shrinkText(self):
        self.changeFontSize(-2)

    def wheelEvent(self, event):
        # Check if Ctrl key is pressed
        if QApplication.keyboardModifiers() == Qt.ControlModifier:
            # Determine whether the scroll is up or down
            delta = event.angleDelta().y()
            if delta > 0:
                self.enlargeText()
            elif delta < 0:
                self.shrinkText()

    def refresh_recheck(self):
        reload_bloom_filter()
        cursor_position = self.text_edit.textCursor()
        text_content = self.text_edit.toHtml()
        plain_text = self.text_edit.toPlainText()
        content_for_bloom = [get_clean_words_for_dictionary(word) for word in plain_text.split() if word if
                             len(word) > 1]
        wrong_words = start_bloom(content_for_bloom)
        # Highlight incorrect words in the editor
        highlighted_content = text_content
        for word in wrong_words:
            highlighted_content = highlighted_content.replace(word, f'<span style="color:red">{word}</span>')
        # Update the editor with the highlighted content
        self.text_edit.setHtml(highlighted_content)
        # Restore the cursor position
        self.text_edit.setTextCursor(cursor_position)
    def eventFilter(self, obj, event):
        if obj == self.text_edit and event.type() == QEvent.KeyPress:
            if event.key() == Qt.Key_Space:
                self.spacebarClicked()
                return True  # Event handled
        return super().eventFilter(obj, event)

    def spacebarClicked(self):
        cursor = self.text_edit.textCursor()
        cursor_position = self.text_edit.textCursor()
        # Save the current position of the cursor
        original_position = cursor.position()
        # Insert a space
        cursor.insertText(" ")
        # Move the cursor to the left of the current position
        cursor.movePosition(QTextCursor.WordLeft)
        # Retrieve the entire word to the left of the cursor using a regular expression
        cursor.movePosition(QTextCursor.EndOfWord, QTextCursor.KeepAnchor)
        word_left_of_cursor = cursor.selectedText()
        # Restore the cursor to the original position
        #cursor.setPosition(original_position)
        if has_letters_or_digits(word_left_of_cursor):
            print("correct word")
        elif not bloom_lookup(word_left_of_cursor):
            wrong_word = f'<span style="color:red">{word_left_of_cursor.strip()}</span> '
            #self.text_edit.setHtml(self.text_edit.toHtml().replace(word_left_of_cursor, wrong_word))
           # self.text_edit.setTextCursor(cursor_position)
           # cursor.setPosition(original_position)
            #self.text_edit.setHtml(wrong_word)
            html_content = self.text_edit.toHtml()
            new_html_content = html_content.replace(word_left_of_cursor.lstrip(), wrong_word.strip(), 1)
            # Set the new HTML content
            self.text_edit.setHtml(new_html_content)
            # Move the cursor to the right of the replaced word
            cursor.movePosition(QTextCursor.Right)
            self.text_edit.setTextCursor(cursor)


    def getWordLeftOfCursor(self):
        cursor = self.text_edit.textCursor()
        # Save the current position of the cursor
        original_position = cursor.position()
        # Move the cursor to the left of the current position
        cursor.movePosition(QTextCursor.WordLeft)
        # Retrieve the entire word to the left of the cursor using a regular expression
        cursor.movePosition(QTextCursor.EndOfWord, QTextCursor.KeepAnchor)
        word_left_of_cursor = cursor.selectedText()
        # Restore the cursor to the original position
        cursor.setPosition(original_position)
        return word_left_of_cursor.strip()


def has_letters_or_digits(word):
    # Define a regular expression pattern
    pattern = re.compile(r'[a-zA-Z\d]')
    # Use re.search to check if the pattern is present in the word
    return bool(re.search(pattern, word))
def main():
    app = QApplication(sys.argv)
    editor = TextEditor()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()
